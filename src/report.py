"""Report assembly and rendering — MeetingReport → markdown / docx.

This module orchestrates the analysis stage and formats the output.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from config import LLM_MODEL, OUTPUT_DIR
from src.analysis import analyze
from src.schemas import MeetingReport, Transcript


# ─── Orchestration ─────────────────────────────────────────────────────────
def build_report(
    transcript: Transcript, prompt_version: str = "v1"
) -> MeetingReport:
    """Run summarization + extraction and bundle into a MeetingReport.

    Uses the fused single-pass ``analyze`` so the transcript is only sent
    to the map model once (instead of once per task).
    """
    summary, action_items = analyze(transcript, prompt_version=prompt_version)
    return MeetingReport(
        transcript=transcript,
        summary=summary,
        action_items=action_items,
        generated_at=datetime.utcnow().isoformat(),
        prompt_version=prompt_version,
        llm_model=LLM_MODEL,
    )


def save_report(report: MeetingReport, stem: str) -> dict[str, Path]:
    """Write both .md and .docx to data/outputs/."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    md_path = OUTPUT_DIR / f"{stem}.md"
    docx_path = OUTPUT_DIR / f"{stem}.docx"
    md_path.write_text(render_markdown(report), encoding="utf-8")
    _render_docx(report, docx_path)
    return {"md": md_path, "docx": docx_path}


# ─── Formatters ────────────────────────────────────────────────────────────
def render_markdown(report: MeetingReport) -> str:
    s = report.summary
    lines: list[str] = [
        "# Meeting Report",
        "",
        f"_Generated: {report.generated_at}  |  Model: {report.llm_model}  |  "
        f"Prompt: {report.prompt_version}_",
        "",
        "## Overview", s.overview, "",
        "## Key Decisions",
    ]
    lines += [f"- {d}" for d in s.key_decisions] or ["- (none)"]
    lines += ["", "## Discussion Points"]
    lines += [f"- {d}" for d in s.discussion_points] or ["- (none)"]
    lines += ["", "## Open Questions"]
    lines += [f"- {d}" for d in s.open_questions] or ["- (none)"]
    lines += ["", "## Action Items", "",
              "| Owner | Task | Due | Confidence |",
              "|-------|------|-----|------------|"]
    for a in report.action_items:
        lines.append(
            f"| {a.owner} | {a.task} | {a.due_date or '—'} | {a.confidence:.2f} |"
        )
    return "\n".join(lines)


def _render_docx(report: MeetingReport, out_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Meeting Report", level=0)
    doc.add_paragraph(
        f"Generated: {report.generated_at} | Model: {report.llm_model} | "
        f"Prompt: {report.prompt_version}"
    )

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(report.summary.overview)

    for title, items in [
        ("Key Decisions", report.summary.key_decisions),
        ("Discussion Points", report.summary.discussion_points),
        ("Open Questions", report.summary.open_questions),
    ]:
        doc.add_heading(title, level=1)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Action Items", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "Owner", "Task", "Due", "Confidence"
    )
    for a in report.action_items:
        row = table.add_row().cells
        row[0].text = a.owner
        row[1].text = a.task
        row[2].text = a.due_date or "—"
        row[3].text = f"{a.confidence:.2f}"

    doc.save(str(out_path))
    return out_path
